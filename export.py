from typing import Dict, List
from jinja2 import Template
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def render_resume_md(data: Dict, jinja_template: str) -> str:
    tmpl = Template(jinja_template)
    return tmpl.render(data)

def render_cover_md(data: Dict, jinja_template: str) -> str:
    tmpl = Template(jinja_template)
    return tmpl.render(data)



def md_to_docx(md_text: str, title: str = "Document") -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def add_bullet(text: str):
        p = doc.add_paragraph(text)
        p.style = "List Bullet"

    lines = md_text.splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")  # blank line
            continue
        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith("- "):
            add_bullet(line[2:].strip())
        elif line.startswith("**") and line.endswith("**"):
            # bold standalone line (e.g., role/company)
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
        else:
            doc.add_paragraph(line)
    return doc


def save_docx(doc: Document, filepath: str):
    doc.save(filepath)



